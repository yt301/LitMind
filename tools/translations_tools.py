import os
from docx import Document
from PyPDF2 import PdfReader




def read_file_content(filepath):
    """统一读取文件内容，支持txt, md, docx, pdf"""
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()

    try:
        if ext in ('.txt', '.md'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.docx':
            doc = Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext == '.pdf':
            text = ""
            with open(filepath, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    except UnicodeDecodeError:
        raise ValueError("文件编码不支持，请确保是UTF-8编码的文本文件")
    except Exception as e:
        raise ValueError(f"读取文件失败: {str(e)}")


async def save_translated_file(content, original_filename, user_dir):
    """保存翻译后的文件"""
    # 生成翻译后的文件名
    name, ext = os.path.splitext(original_filename)
    translated_filename = f"{name}_translated{ext}"
    translated_filepath = os.path.join(user_dir, translated_filename)

    # 根据原文件类型保存
    if ext.lower() == '.docx':
        doc = Document()
        doc.add_paragraph(content)
        doc.save(translated_filepath)
    elif ext.lower() == '.pdf':
        # PDF保存需要更复杂的处理，这里简化为文本文件
        with open(translated_filepath, 'w', encoding='utf-8') as f:
            f.write(content)
    else:  # txt, md等文本文件
        with open(translated_filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    return translated_filename, translated_filepath

def detect_file_type(filepath):
    """通过文件扩展名检测文件类型"""
    ext = os.path.splitext(filepath)[1].lower()
    type_map = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.pdf': 'application/pdf'
    }
    return type_map.get(ext, 'application/octet-stream')